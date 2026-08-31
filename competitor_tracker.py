"""
competitor_tracker.py
Weekly competitor update tracker.

Checks:
  - iOS App Store version/release notes (via official iTunes Lookup API)
  - Google Play version/release notes (via google-play-scraper library)
  - RSS feeds (official blogs/newsrooms)

Writes a Word report to ./digests/digest-<date>.docx (and a Markdown
copy alongside it) containing only what's NEW since the last run
(tracked in state.json).
"""

import json
import os
from datetime import datetime, timedelta

import feedparser
import requests
from google_play_scraper import app as gp_app
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------- CONFIG: edit these ----------------

# App Store numeric IDs (from the app's App Store URL, e.g.
# apps.apple.com/us/app/tiktok/id835599320 -> 835599320)
ITUNES_APPS = {
    "TikTok": 835599320,
    "Bigo Live": 1077137248,
}

# Google Play package IDs (from the Play Store URL's id= param)
PLAY_STORE_APPS = {
    "TikTok": "com.zhiliaoapp.musically",
    "Bigo Live": "sg.bigo.live",
}

# Official blogs / newsrooms. TikTok does not publish a native RSS feed,
# so these use feed URLs generated at https://rss.app (free plan) pointed
# at each company's newsroom/blog page. Paste your generated URLs below,
# replacing the "PASTE_..." placeholders. Leave a line commented out (#)
# to skip that source until you have a working feed URL for it.
RSS_FEEDS = {
    # "TikTok Newsroom": "PASTE_YOUR_RSS_APP_URL_HERE",  # skipped - JS-heavy, no working feed found
    "TikTok Business Blog": "https://rss.app/feeds/XkhI9YJuLjt29E08.xml",
    "Bigo Events": "https://rss.app/feeds/3jCaB1MUpFXQUK4I.xml",
}

STATE_FILE = "state.json"
OUTPUT_DIR = "digests"
RSS_LOOKBACK_DAYS = 8

# ------------------------------------------------------


def load_state():
    if os.path.exists(STATE_FILE):
        with open(STATE_FILE) as f:
            return json.load(f)
    return {}


def save_state(state):
    with open(STATE_FILE, "w") as f:
        json.dump(state, f, indent=2)


def check_itunes(state, updates):
    for name, app_id in ITUNES_APPS.items():
        try:
            r = requests.get(
                f"https://itunes.apple.com/lookup?id={app_id}", timeout=15
            )
            r.raise_for_status()
            results = r.json().get("results", [])
            if not results:
                continue
            data = results[0]
            version = data.get("version")
            notes = (data.get("releaseNotes") or "").strip()
            key = f"ios:{name}"
            if state.get(key) != version:
                updates.append(
                    {
                        "app": name,
                        "source": "App Store",
                        "type": "Version update",
                        "detail": f"v{version} — {notes}" if notes else f"v{version}",
                    }
                )
                state[key] = version
        except Exception as e:
            updates.append(
                {"app": name, "source": "App Store", "type": "ERROR", "detail": str(e)}
            )


def check_play_store(state, updates):
    for name, pkg in PLAY_STORE_APPS.items():
        try:
            data = gp_app(pkg)
            version = data.get("version")
            notes = (data.get("recentChanges") or "").strip()
            key = f"android:{name}"
            if state.get(key) != version:
                updates.append(
                    {
                        "app": name,
                        "source": "Google Play",
                        "type": "Version update",
                        "detail": f"v{version} — {notes}" if notes else f"v{version}",
                    }
                )
                state[key] = version
        except Exception as e:
            updates.append(
                {"app": name, "source": "Google Play", "type": "ERROR", "detail": str(e)}
            )


def check_rss(state, updates):
    for name, url in RSS_FEEDS.items():
        if not url or url.startswith("PASTE_"):
            continue  # skip until a real feed URL is filled in
        try:
            feed = feedparser.parse(url)
            seen = set(state.get(f"rss:{name}", []))
            new_seen = list(seen)
            for entry in feed.entries[:20]:
                link = entry.get("link")
                if link and link not in seen:
                    updates.append(
                        {
                            "app": name,
                            "source": name,
                            "type": "News/Blog post",
                            "detail": f"{entry.get('title', 'Untitled')} — {link}",
                        }
                    )
                    new_seen.append(link)
            state[f"rss:{name}"] = new_seen[-200:]
        except Exception as e:
            updates.append({"app": name, "source": name, "type": "ERROR", "detail": str(e)})


def group_by_app(updates):
    by_app = {}
    for u in updates:
        by_app.setdefault(u["app"], []).append(u)
    return by_app


def write_digest_md(updates, date_str):
    """Plain-text Markdown copy — handy for quick viewing on GitHub itself."""
    path = os.path.join(OUTPUT_DIR, f"digest-{date_str}.md")
    with open(path, "w") as f:
        f.write(f"# Competitor Update Digest — {date_str}\n\n")
        if not updates:
            f.write("No new updates detected since the last run.\n")
        else:
            for app_name, items in group_by_app(updates).items():
                f.write(f"## {app_name}\n\n")
                for u in items:
                    f.write(f"- **[{u['source']}] {u['type']}:** {u['detail']}\n")
                f.write("\n")
    return path


def write_digest_docx(updates, date_str):
    """The Word report — this is the file to open and share with the team."""
    path = os.path.join(OUTPUT_DIR, f"digest-{date_str}.docx")
    doc = Document()

    title = doc.add_heading(f"Competitor Update Digest — {date_str}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if not updates:
        p = doc.add_paragraph("No new updates detected since the last run.")
        p.runs[0].italic = True
    else:
        for app_name, items in group_by_app(updates).items():
            doc.add_heading(app_name, level=1)
            for u in items:
                p = doc.add_paragraph(style="List Bullet")
                tag_run = p.add_run(f"[{u['source']}] ")
                tag_run.bold = True
                tag_run.font.size = Pt(10)
                if u["type"] == "ERROR":
                    tag_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                type_run = p.add_run(f"{u['type']}: ")
                type_run.bold = True
                p.add_run(u["detail"])

    doc.save(path)
    return path


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    date_str = datetime.utcnow().strftime("%Y-%m-%d")

    state = load_state()
    updates = []
    check_itunes(state, updates)
    check_play_store(state, updates)
    check_rss(state, updates)

    docx_path = write_digest_docx(updates, date_str)
    md_path = write_digest_md(updates, date_str)
    save_state(state)
    print(f"Digest written to {docx_path} and {md_path} ({len(updates)} item(s))")


if __name__ == "__main__":
    main()
